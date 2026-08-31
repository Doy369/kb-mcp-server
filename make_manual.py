"""生成《企业客服知识库桌面客户端 · 操作手册》Word 文档。

运行: python make_manual.py  ->  产出 操作手册.docx (项目根目录)
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

FONT = "Microsoft YaHei"


def set_cjk(style, font_name=FONT):
    """让标题/正文等样式正确显示中文（设置 eastAsia 字体）。"""
    style.font.name = font_name
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)


def add_steps(doc, items):
    for s in items:
        doc.add_paragraph(s, style="List Number")


def add_bullets(doc, items):
    for s in items:
        doc.add_paragraph(s, style="List Bullet")


def main():
    doc = Document()

    # 默认正文字体
    normal = doc.styles["Normal"]
    normal.font.size = Pt(11)
    set_cjk(normal)
    for name in ("Heading 1", "Heading 2", "Title"):
        try:
            set_cjk(doc.styles[name])
        except KeyError:
            pass

    # 标题
    t = doc.add_heading("企业客服知识库桌面客户端", level=0)
    sub = doc.add_paragraph("操作手册  ·  适用对象：客服 / 文职人员")
    sub.runs[0].italic = True
    sub.runs[0].font.size = Pt(12)
    meta = doc.add_paragraph("版本：v1.0　|　最后更新：2026-08　|　程序：kb-mcp-client.exe")
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph(
        "本手册面向客服与文职人员，介绍如何在不需要安装 Python、不需要敲命令的情况下，"
        "通过桌面客户端使用企业客服知识库系统。双击程序即可在本机弹出一个窗口，"
        "进行知识库管理、智能问答与接口配置。"
    )

    # 一、软件简介
    doc.add_heading("一、软件简介", level=1)
    add_bullets(doc, [
        "这是一个“桌面客户端”（不是网页）。双击后自动弹出窗口，无需打开浏览器。",
        "后端在本机运行，数据保存在你电脑的“用户目录”下，不外传到任何服务器。",
        "核心能力：管理知识库文档、基于知识库回答问题（RAG 检索增强）、配置外部业务接口（如订单状态、DeepSeek 大模型）。",
    ])

    # 二、运行环境要求
    doc.add_heading("二、运行环境要求", level=1)
    add_bullets(doc, [
        "操作系统：Windows 10 或 Windows 11（64 位）。",
        "必须组件：Microsoft WebView2 运行库（Win10/11 通常已预装；若打开后窗口空白，请到微软官网下载安装 “WebView2 Runtime”）。",
        "无需安装 Python，日常使用无需联网（除非要调用外部接口 / 大模型）。",
    ])

    # 三、安装与首次启动
    doc.add_heading("三、安装与首次启动", level=1)
    add_steps(doc, [
        "把 kb-mcp-client.exe（建议使用已修复 DPI 的 v2 版本）放到任意普通文件夹。建议不要放在 C:\\Program Files，以免没有写入权限。",
        "双击 kb-mcp-client.exe。",
        "首次启动会自动完成三件事：① 在你电脑的“应用数据”目录创建数据文件夹；② 预置 6 篇示例知识库文档；③ 弹出客户端窗口。",
        "以后每次双击即可直接使用，你的配置与新增文档都会自动保留。",
    ])

    # 四、界面总览
    doc.add_heading("四、界面总览", level=1)
    doc.add_paragraph("窗口分为左右两区：")
    add_bullets(doc, [
        "左侧（主区）：知识库管理、检索结果、接口配置等标签页。",
        "右侧（聊天区，约占 1/3 宽）：直接输入问题，获得基于知识库的回答。",
        "顶部标题栏显示“企业客服知识库控制台”。",
    ])

    # 五、知识库管理
    doc.add_heading("五、知识库管理", level=1)
    add_steps(doc, [
        "进入“知识库”或“文档”标签页。",
        "点击“添加文档 / 上传”，选择要加入的文档（支持文本、常见问题等）。",
        "系统会自动切分、向量化并建立索引，稍候即可检索。",
        "可在文档列表中查看、删除文档。",
    ])
    doc.add_paragraph(
        "提示：系统预置的示例文档包括《退货退款政策》《物流配送 FAQ》《服务级别协议 SLA》"
        "《账户与计费》《客服常见问题》《保修条款》，可参考其写法整理你自己的知识库。"
    )

    # 六、智能问答
    doc.add_heading("六、智能问答", level=1)
    add_steps(doc, [
        "在右侧聊天框输入客户问题，例如“退货后多久能收到退款？”。",
        "回车发送，系统会检索知识库并生成答案，同时标注引用来源与置信度（高 / 中 / 低）。",
        "若置信度为“低”，代表知识库中缺少相关内容，建议补充文档或由人工回复。",
    ])

    # 七、接口配置（文职向导）
    doc.add_heading("七、接口配置（文职向导）", level=1)
    doc.add_paragraph("打开“接口配置”标签页，按提示填写：")
    add_bullets(doc, [
        "演示模式：新手可先开启，用内置示例数据体验，无需任何密钥。",
        "一键模型模板：点击后自动填入常用大模型（如 DeepSeek）的配置模板，只需补上你的 API Key。",
        "外部业务接口（如订单状态）：填写接口地址，系统会“自动识别字段”并以下拉方式让你勾选对应关系"
        "（例如 orderState = 订单状态、logistics.company = 物流公司），无需手写代码。",
        "密钥安全：在面板填写的密钥只保存在你电脑本地文件，不会上传。",
    ])

    # 八、常见问题 FAQ
    doc.add_heading("八、常见问题（FAQ）", level=1)
    rows = [
        ("双击后窗口空白 / 只有右侧一条窄缝？",
         "多为缺少 WebView2 运行库或缩放显示异常。请安装 WebView2 Runtime；"
         "已发布的 v2 版本已内置 DPI 修复，使用 v2 即可避免此问题。"),
        ("程序打不开 / 闪退？",
         "确认未放在只读目录（如 C:\\Program Files）；用普通文件夹重试。"
         "也可能是旧实例仍在运行：任务管理器结束 kb-mcp-client.exe 后重新双击。"),
        ("想换一个端口？",
         "高级用法：设置环境变量 KB_WEB_PORT=9000 后再启动程序即可。"),
        ("问答显示“知识库中没有相关内容”？",
         "先到“知识库管理”添加相关文档，并等待索引建立完成后再问。"),
        ("配置或文档丢失了？",
         "数据保存在 C:\\Users\\你的用户名\\AppData\\Roaming\\kb-mcp-server\\，请勿删除该目录。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "问题"
    hdr[1].text = "解决办法"
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.bold = True
                set_cjk(doc.styles["Normal"])  # 保证表内中文
    for q, a in rows:
        cells = table.add_row().cells
        cells[0].text = q
        cells[1].text = a

    # 九、数据备份与卸载
    doc.add_heading("九、数据备份与卸载", level=1)
    add_bullets(doc, [
        "备份：复制 C:\\Users\\你的用户名\\AppData\\Roaming\\kb-mcp-server\\ 整个文件夹即可。",
        "卸载：直接删除 exe 程序文件即可；如需彻底清除数据，一并删除上述 AppData 目录。",
    ])

    # 十、安全须知
    doc.add_heading("十、安全须知", level=1)
    add_bullets(doc, [
        "本系统为本地客户端，知识库与密钥均存储在你电脑本地，默认不上传任何服务器。",
        "DeepSeek 等大模型接口，仅在你配置了 Key 并提问时，才会把问题文本发往对应服务商。",
        "请勿在公共电脑上保存含敏感业务数据的知识库；如必须，请用完即清理 AppData 目录。",
    ])

    out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "操作手册.docx")
    doc.save(out)
    print("saved ->", out)


if __name__ == "__main__":
    main()
